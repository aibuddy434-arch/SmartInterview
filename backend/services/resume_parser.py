import logging
import os
from typing import Optional
import PyPDF2
from docx import Document

logger = logging.getLogger(__name__)

class ResumeParserService:
    """Service for parsing resume files (PDF, DOCX, DOC, TXT) to extract text content."""
    
    @staticmethod
    async def parse_resume(file_path: str) -> str:
        """
        Parse a resume file and extract text content.
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            Extracted text content as a string. Returns empty string if parsing fails.
        """
        if not file_path or not os.path.exists(file_path):
            logger.warning(f"Resume file not found: {file_path}")
            return ""
        
        try:
            # Detect file type based on extension
            file_extension = os.path.splitext(file_path)[1].lower()
            
            if file_extension == '.pdf':
                return await ResumeParserService._parse_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                return await ResumeParserService._parse_docx(file_path)
            elif file_extension == '.txt':
                return await ResumeParserService._parse_txt(file_path)
            else:
                logger.warning(f"Unsupported file format: {file_extension}")
                return ""
                
        except Exception as e:
            logger.error(f"Failed to parse resume {file_path}: {str(e)}")
            return ""
    
    @staticmethod
    async def _parse_pdf(file_path: str) -> str:
        """Extract text from PDF file."""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            logger.info(f"Successfully parsed PDF: {file_path} ({len(text)} characters)")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {str(e)}")
            return ""
    
    @staticmethod
    async def _parse_docx(file_path: str) -> str:
        """Extract text from DOCX/DOC file."""
        try:
            doc = Document(file_path)
            text = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
            
            result = "\n".join(text)
            logger.info(f"Successfully parsed DOCX: {file_path} ({len(result)} characters)")
            return result.strip()
            
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {str(e)}")
            return ""
    
    @staticmethod
    async def _parse_txt(file_path: str) -> str:
        """Extract text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            logger.info(f"Successfully parsed TXT: {file_path} ({len(text)} characters)")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error parsing TXT {file_path}: {str(e)}")
            return ""

# Global service instance
resume_parser_service = ResumeParserService()


